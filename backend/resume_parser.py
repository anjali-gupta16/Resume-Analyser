"""
resume_parser.py — Extract plain text from PDF and DOCX files
"""
import io
import re
from pathlib import Path


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pypdf (much faster than pdfminer)."""
    from pypdf import PdfReader
    try:
        reader = PdfReader(file_path)
        text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
        return clean_text("\n".join(text_parts))
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document
    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        full_text = "\n".join(paragraphs)
        return clean_text(full_text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")


def clean_text(text: str) -> str:
    """Normalise whitespace and remove non-printable characters."""
    if not text:
        return ""
    # Replace multiple newlines with double newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Replace multiple spaces with single space
    text = re.sub(r' {2,}', ' ', text)
    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    # Filter out empty lines sequences
    cleaned = "\n".join(line for line in lines if line or lines.index(line) == 0)
    return cleaned.strip()


def parse_resume(file_path: str) -> str:
    """Auto-detect file type and extract text."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload PDF or DOCX.")
